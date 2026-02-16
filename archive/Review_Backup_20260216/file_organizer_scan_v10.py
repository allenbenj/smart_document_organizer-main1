import os
import hashlib
import json
import sqlite3
from datetime import datetime
import concurrent.futures
from typing import List, Dict, Any, Set, Iterator
import logging
import traceback
import time
import docx
import PyPDF2
from PIL import Image
import imagehash
import nltk
from nltk.tokenize import word_tokenize
from nltk.corpus import stopwords
from langdetect import detect
import cv2
import numpy as np
import filetype
import pickle
import pandas as pd
import asyncio

# Set up logging
logging.basicConfig(level=logging.DEBUG, format='%(asctime)s - %(levelname)s - %(message)s')
logger = logging.getLogger(__name__)

# Download necessary NLTK data
nltk.download('punkt', quiet=True)
nltk.download('stopwords', quiet=True)

# Constants
NUM_THREADS = max(12, os.cpu_count() or 1)
BATCH_SIZE = 1000
CHECKPOINT_FILE = 'file_organizer_checkpoint.pkl'

def retry_operation(func, max_attempts=3, delay=1):
    def wrapper(*args, **kwargs):
        attempts = 0
        while attempts < max_attempts:
            try:
                return func(*args, **kwargs)
            except Exception as e:
                attempts += 1
                if attempts == max_attempts:
                    raise
                logger.warning(f"Operation failed. Retrying in {delay} seconds. Error: {str(e)}")
                time.sleep(delay)
    return wrapper

@retry_operation
def hash_file(filepath: str) -> str:
    logger.debug(f"Hashing file: {filepath}")
    blake2b_hash = hashlib.blake2b()
    with open(filepath, "rb") as f:
        for byte_block in iter(lambda: f.read(4096), b""):
            blake2b_hash.update(byte_block)
    return blake2b_hash.hexdigest()

@retry_operation
def get_file_metadata(filepath: str) -> Dict[str, Any]:
    logger.debug(f"Getting metadata for file: {filepath}")
    stat = os.stat(filepath)
    return {
        "size": stat.st_size,
        "created": datetime.fromtimestamp(stat.st_ctime).isoformat(),
        "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
    }

def extract_keywords(text: str, max_keywords=10) -> List[str]:
    stop_words = set(stopwords.words('english'))
    words = word_tokenize(text.lower())
    keywords = [word for word in words if word.isalnum() and word not in stop_words]
    return keywords[:max_keywords]

def detect_language(text: str) -> str:
    try:
        return detect(text)
    except:
        return "unknown"

def analyze_image(filepath: str) -> Dict[str, Any]:
    with Image.open(filepath) as img:
        # Basic image info
        info = {
            "format": img.format,
            "mode": img.mode,
            "size": img.size,
        }
        
        # Calculate image hash for similarity comparison
        info["hash"] = str(imagehash.average_hash(img))
        
        # Convert to OpenCV format for additional analysis
        cv_image = cv2.cvtColor(np.array(img), cv2.COLOR_RGB2BGR)
        
        # Detect faces
        face_cascade = cv2.CascadeClassifier(cv2.data.haarcascades + 'haarcascade_frontalface_default.xml')
        gray = cv2.cvtColor(cv_image, cv2.COLOR_BGR2GRAY)
        faces = face_cascade.detectMultiScale(gray, 1.1, 4)
        info["faces_detected"] = len(faces)
        
        # Determine if image is likely a photo or graphic
        edges = cv2.Canny(cv_image, 100, 200)
        edge_ratio = np.sum(edges > 0) / (edges.shape[0] * edges.shape[1])
        info["likely_photo"] = edge_ratio < 0.1  # Threshold can be adjusted
        
        return info

def get_mime_type(filepath: str) -> str:
    try:
        mime_type = filetype.guess(filepath)
        if mime_type:
            return mime_type.mime
    except Exception as e:
        logger.warning(f"FileType failed to detect file type for {filepath}: {str(e)}")
    
    return "application/octet-stream"  # Default to binary data if detection fails

@retry_operation
def extract_file_sample(filepath: str, max_size: int = 1024) -> Dict[str, Any]:
    logger.debug(f"Extracting sample from file: {filepath}")
    mime_type = get_mime_type(filepath)
    result = {"mime_type": mime_type, "sample": "", "analysis": {}}
    
    try:
        if mime_type.startswith('text/'):
            with open(filepath, 'r', errors='ignore') as f:
                text = f.read(max_size)
                result["sample"] = text
                result["analysis"]["keywords"] = extract_keywords(text)
                result["analysis"]["language"] = detect_language(text)
        
        elif mime_type == 'application/pdf':
            with open(filepath, 'rb') as f:
                pdf_reader = PyPDF2.PdfReader(f)
                if pdf_reader.pages:
                    text = ' '.join(page.extract_text()[:max_size] for page in pdf_reader.pages[:2])
                    result["sample"] = text
                    result["analysis"]["keywords"] = extract_keywords(text)
                    result["analysis"]["language"] = detect_language(text)
                else:
                    result["sample"] = "EMPTY_PDF"

        elif mime_type == 'application/vnd.openxmlformats-officedocument.wordprocessingml.document':
            doc = docx.Document(filepath)
            if doc.paragraphs:
                text = ' '.join(paragraph.text for paragraph in doc.paragraphs)[:max_size]
                result["sample"] = text
                result["analysis"]["keywords"] = extract_keywords(text)
                result["analysis"]["language"] = detect_language(text)
            else:
                result["sample"] = "EMPTY_DOCX"
        
        elif mime_type.startswith('image/'):
            result["analysis"] = analyze_image(filepath)
            result["sample"] = f"Image: {result['analysis']['format']}, {result['analysis']['size'][0]}x{result['analysis']['size'][1]}, Mode: {result['analysis']['mode']}"
        
        elif mime_type.startswith('audio/'):
            result["sample"] = f"Audio file: {mime_type}"
        
        elif mime_type.startswith('video/'):
            result["sample"] = f"Video file: {mime_type}"
        
        else:
            result["sample"] = f"Unsupported file type: {mime_type}"
    
    except Exception as e:
        logger.error(f"Error extracting sample from file {filepath}: {str(e)}")
        result["sample"] = "ERROR_EXTRACTING_SAMPLE"
        result["error"] = str(e)
    
    return result

def process_file(filepath: str) -> Dict[str, Any]:
    logger.info(f"Processing file: {filepath}")
    try:
        file_info = {
            "path": filepath,
            "name": os.path.basename(filepath),
            "hash": hash_file(filepath),
            "metadata": get_file_metadata(filepath),
        }
        sample_info = extract_file_sample(filepath)
        file_info.update(sample_info)
        return file_info
    except Exception as e:
        logger.error(f"Error processing file {filepath}: {str(e)}")
        logger.error(traceback.format_exc())
        return {"path": filepath, "error": str(e)}

def scan_directory(folder_path: str, excluded_types: Set[str] = set(), processed_files: Set[str] = set()) -> Iterator[Dict[str, Any]]:
    logger.info(f"Starting directory scan: {folder_path}")
    if not os.path.exists(folder_path):
        logger.error(f"Directory does not exist: {folder_path}")
        return iter([])

    total_files = sum([len(files) for _, _, files in os.walk(folder_path)])
    processed_count = 0

    with concurrent.futures.ThreadPoolExecutor(max_workers=NUM_THREADS) as executor:
        futures = []
        for dirpath, _, filenames in os.walk(folder_path):
            for filename in filenames:
                filepath = os.path.join(dirpath, filename)
                if filepath not in processed_files:
                    futures.append(executor.submit(process_file, filepath))

        for future in concurrent.futures.as_completed(futures):
            try:
                result = future.result()
                if result and result.get('mime_type') not in excluded_types:
                    yield result
                processed_count += 1
                if processed_count % 100 == 0:
                    logger.info(f"Processed {processed_count}/{total_files} files")
            except Exception as e:
                logger.error(f"Error processing a file: {str(e)}")
                logger.error(traceback.format_exc())

    logger.info(f"Completed directory scan. Processed {processed_count}/{total_files} files.")

def create_sqlite_db(db_name: str):
    conn = sqlite3.connect(db_name)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS files
                 (path TEXT PRIMARY KEY, name TEXT, hash TEXT, size INTEGER, 
                  created TEXT, modified TEXT, mime_type TEXT, sample TEXT)''')
    c.execute('''CREATE TABLE IF NOT EXISTS file_analysis
                 (file_path TEXT PRIMARY KEY, keywords TEXT, language TEXT, 
                  image_format TEXT, image_mode TEXT, image_size TEXT, 
                  image_hash TEXT, faces_detected INTEGER, likely_photo INTEGER)''')
    conn.commit()
    return conn

def insert_file_data(conn, file_data: Dict[str, Any]):
    c = conn.cursor()
    c.execute('''INSERT OR REPLACE INTO files VALUES (?, ?, ?, ?, ?, ?, ?, ?)''',
              (file_data['path'], file_data['name'], file_data['hash'],
               file_data['metadata']['size'], file_data['metadata']['created'],
               file_data['metadata']['modified'], file_data.get('mime_type', ''),
               file_data.get('sample', '')))
    
    analysis = file_data.get('analysis', {})
    if analysis:
        c.execute('''INSERT OR REPLACE INTO file_analysis VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?)''',
                  (file_data['path'], ','.join(analysis.get('keywords', [])),
                   analysis.get('language', ''), analysis.get('format', ''),
                   analysis.get('mode', ''), str(analysis.get('size', '')),
                   analysis.get('hash', ''), analysis.get('faces_detected', 0),
                   1 if analysis.get('likely_photo', False) else 0))
    conn.commit()

def process_batch(batch: List[Dict[str, Any]], json_file, db_conn):
    # Write to JSON
    for item in batch:
        json.dump(item, json_file)
        json_file.write('\n')
    
    # Write to SQLite
    for item in batch:
        insert_file_data(db_conn, item)

def save_checkpoint(processed_files: Set[str]):
    with open(CHECKPOINT_FILE, 'wb') as f:
        pickle.dump(processed_files, f)

def load_checkpoint() -> Set[str]:
    if os.path.exists(CHECKPOINT_FILE):
        with open(CHECKPOINT_FILE, 'rb') as f:
            return pickle.load(f)
    return set()

def main():
    folder_path = input("Enter the path of the folder to scan (including all subfolders): ")
    excluded_types = set()  # Add MIME types to exclude, e.g., {'image/jpeg', 'audio/mpeg'}
    logger.info(f"Starting file organization process for folder and subfolders: {folder_path}")
    
    processed_files = load_checkpoint()
    logger.info(f"Loaded {len(processed_files)} processed files from checkpoint")

    file_data_iterator = scan_directory(folder_path, excluded_types, processed_files)
    
    json_file = open('file_data.jsonl', 'a')  # Append mode to continue from last run
    db_conn = create_sqlite_db('file_data.db')
    
    batch = []
    try:
        for item in file_data_iterator:
            batch.append(item)
            if len(batch) >= BATCH_SIZE:
                process_batch(batch, json_file, db_conn)
                processed_files.update(file['path'] for file in batch)
                save_checkpoint(processed_files)
                batch = []
        
        if batch:  # Process any remaining items
            process_batch(batch, json_file, db_conn)
            processed_files.update(file['path'] for file in batch)
            save_checkpoint(processed_files)
    
    except Exception as e:
        logger.error(f"An error occurred during processing: {str(e)}")
        logger.error(traceback.format_exc())
    finally:
        json_file.close()
        db_conn.close()
        save_checkpoint(processed_files)
    
    logger.info("File organization process completed or interrupted.")
    print("Data saved to file_data.jsonl and file_data.db")
    print(f"Processed files checkpoint saved to {CHECKPOINT_FILE}")

if __name__ == "__main__":
    main()
